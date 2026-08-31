"""
Multi-Format Document Exporter for GenDoc AI.
Generates styled PDF (via ReportLab), DOCX (via python-docx), and Markdown.
"""

import os
import re
from typing import Dict, Any, List
from io import BytesIO

# ReportLab imports for PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted, PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas

# python-docx imports
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class NumberedCanvas(canvas.Canvas):
    """Canvas that performs a two-pass calculation to draw 'Page X of Y' in footer."""
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super(NumberedCanvas, self).showPage()
        super(NumberedCanvas, self).save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#64748b"))

        # Don't draw header/footer on cover page if page 1
        if self._pageNumber > 1:
            # Top Header
            self.drawString(54, 750, "GenDoc AI – Technical Documentation")
            self.setStrokeColor(colors.HexColor("#e2e8f0"))
            self.setLineWidth(0.5)
            self.line(54, 742, 558, 742)

            # Bottom Footer
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(558, 40, page_text)
            self.drawString(54, 40, "Confidential & Proprietary • Powered by GenDoc AI")
            self.line(54, 52, 558, 52)

        self.restoreState()


class DocExporter:
    """Exports technical documentation to PDF, DOCX, and Markdown formats."""

    @classmethod
    def generate_pdf(cls, title: str, markdown_content: str, output_path: str = None) -> bytes:
        """Generates a professional, styled PDF with cover page, headers, tables, and code blocks."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer if output_path is None else output_path,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()

        # Custom Styles
        title_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=26,
            leading=32,
            textColor=colors.HexColor('#0f172a'),
            alignment=0,
            spaceAfter=15
        )

        subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=13,
            leading=18,
            textColor=colors.HexColor('#475569'),
            alignment=0,
            spaceAfter=25
        )

        h1_style = ParagraphStyle(
            'SectionH1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#1e293b'),
            spaceBefore=18,
            spaceAfter=10,
            keepWithNext=True
        )

        h2_style = ParagraphStyle(
            'SectionH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#334155'),
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True
        )

        h3_style = ParagraphStyle(
            'SectionH3',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=15,
            textColor=colors.HexColor('#475569'),
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'Body',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=8
        )

        code_style = ParagraphStyle(
            'CodeStyle',
            fontName='Courier',
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor('#0f172a')
        )

        story = []

        # Cover Banner Badge
        badge_table = Table([["⚡ POWERED BY GENDOC AI"]], colWidths=[180])
        badge_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#e0e7ff')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#4338ca')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(badge_table)
        story.append(Spacer(1, 20))

        # Title and Subtitle
        clean_title = title.replace("#", "").strip()
        story.append(Paragraph(clean_title, title_style))
        story.append(Paragraph("Comprehensive Technical Specification & Architecture Manual", subtitle_style))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#6366f1"), spaceAfter=20))

        # Parse Markdown lines into ReportLab flowables
        lines = markdown_content.split("\n")
        in_code_block = False
        code_buffer = []

        for line in lines:
            line_str = line.rstrip()

            # Code Block Start/End
            if line_str.startswith("```"):
                if in_code_block:
                    # End of code block
                    code_text = "\n".join(code_buffer)
                    code_table = Table([[Preformatted(code_text, code_style)]], colWidths=[500])
                    code_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
                        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
                        ('PADDING', (0, 0), (-1, -1), 8),
                    ]))
                    story.append(code_table)
                    story.append(Spacer(1, 10))
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_buffer = []
                continue

            if in_code_block:
                code_buffer.append(line_str)
                continue

            # Heading 1
            if line_str.startswith("# "):
                story.append(Paragraph(line_str[2:], h1_style))
            # Heading 2
            elif line_str.startswith("## "):
                story.append(Paragraph(line_str[3:], h2_style))
            # Heading 3
            elif line_str.startswith("### "):
                story.append(Paragraph(line_str[4:], h3_style))
            # Table Row (Simple Markdown Table Parsing)
            elif line_str.startswith("|") and line_str.endswith("|") and not ("---" in line_str):
                cells = [c.strip() for c in line_str.split("|")[1:-1]]
                row_table = Table([cells], colWidths=[500 / max(len(cells), 1)] * len(cells))
                row_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f1f5f9')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1e293b')),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold' if story and isinstance(story[-1], Paragraph) else 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8.5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                    ('PADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(row_table)
            # Divider
            elif line_str.startswith("---"):
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0"), spaceBefore=10, spaceAfter=10))
            # Bullet point
            elif line_str.startswith("- ") or line_str.startswith("* "):
                bullet_text = f"• {line_str[2:]}"
                story.append(Paragraph(bullet_text, body_style))
            # Regular text
            elif line_str.strip():
                # Escape XML special chars for ReportLab Paragraph
                escaped = line_str.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                # Basic bold rendering: **text** -> <b>text</b>
                escaped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', escaped)
                story.append(Paragraph(escaped, body_style))
            else:
                story.append(Spacer(1, 4))

        doc.build(story, canvasmaker=NumberedCanvas)

        if output_path is None:
            return buffer.getvalue()
        return b""

    @classmethod
    def generate_docx(cls, title: str, markdown_content: str, output_path: str = None) -> bytes:
        """Generates a professional DOCX document with executive styling."""
        doc = docx.Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title.replace("#", "").strip())
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)

        # Subtitle
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run("GenDoc AI Generated Technical Specification & System Architecture")
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(99, 102, 241)

        doc.add_paragraph("-" * 60)

        lines = markdown_content.split("\n")
        in_code = False
        code_lines = []

        for line in lines:
            line_str = line.rstrip()

            if line_str.startswith("```"):
                if in_code:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run("\n".join(code_lines))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(30, 41, 59)
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                    code_lines = []
                continue

            if in_code:
                code_lines.append(line_str)
                continue

            if line_str.startswith("# "):
                h = doc.add_heading(line_str[2:], level=1)
                h.style.font.color.rgb = RGBColor(30, 41, 59)
            elif line_str.startswith("## "):
                h = doc.add_heading(line_str[3:], level=2)
                h.style.font.color.rgb = RGBColor(71, 85, 105)
            elif line_str.startswith("### "):
                h = doc.add_heading(line_str[4:], level=3)
                h.style.font.color.rgb = RGBColor(100, 116, 139)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                p = doc.add_paragraph(line_str[2:], style='List Bullet')
                p.style.font.size = Pt(10)
            elif line_str.strip():
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line_str)
                p = doc.add_paragraph(clean_text)
                p.style.font.size = Pt(10)

        buffer = BytesIO()
        if output_path:
            doc.save(output_path)
            return b""
        else:
            doc.save(buffer)
            return buffer.getvalue()

    @classmethod
    def generate_markdown(cls, markdown_content: str) -> str:
        """Returns clean markdown string."""
        return markdown_content.strip()
